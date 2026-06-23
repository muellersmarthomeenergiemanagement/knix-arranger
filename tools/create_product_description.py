"""
Erstellt eine Produktbeschreibung für KNiX Arranger als Word-Dokument (.docx).

Ausführen:
    python tools/create_product_description.py
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = "KNiX_Arranger_Produktbeschreibung.docx"

C_GREEN_DARK = RGBColor(0x1B, 0x5E, 0x20)
C_GREEN      = RGBColor(0x2E, 0x7D, 0x32)
C_GREY       = RGBColor(0x55, 0x55, 0x55)


def add_heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    run.font.color.rgb = C_GREEN_DARK


def add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    run.font.color.rgb = C_GREEN


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(11)


def build() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Titelseite ────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KNiX Arranger")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = C_GREEN_DARK

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software für die professionelle KNX-Projektplanung")
    run.font.size = Pt(14)
    run.font.color.rgb = C_GREY

    doc.add_paragraph()

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Müller SmartHome & EnergieManagement").font.size = Pt(11)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("info@muellersmarthomeenergiemanagement.ch")
    run.font.size = Pt(11)
    run.font.color.rgb = C_GREEN

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 1 Einleitung ──────────────────────────────────────────────────────
    add_heading1(doc, "Was ist KNiX Arranger?")
    add_body(doc,
        "KNiX Arranger ist eine Desktop-Anwendung für Windows, die den gesamten "
        "Planungsprozess für KNX-Gebäudeautomationsprojekte abbildet – von der "
        "ersten Skizze der Gebäudestruktur bis zur fertigen, ETS6-kompatiblen "
        "Projektdatei samt vollständiger Revisionsdokumentation. Statt "
        "Gruppenadressen, Topologien, Gerätelisten und Berichte von Hand "
        "zusammenzustellen, leitet KNiX Arranger diese Informationen automatisiert "
        "aus den Eingaben des Planers ab – normgerecht nach den Vorgaben der "
        "KNX-Topologie-Regeln und den KNX Swiss Projektrichtlinien.")
    add_body(doc,
        "Das Programm begleitet ein Projekt durch einen klar geführten, "
        "13-stufigen Planungs-Wizard. Jeder Schritt baut auf den Ergebnissen "
        "des vorherigen auf: Aus der Gebäudestruktur ergeben sich die Gewerke "
        "(Licht, Jalousie, Heizung, Lüftung, Energie, Multimedia usw.), aus den "
        "Gewerken die benötigten Aktoren, aus der Aktoranzahl die KNX-Topologie "
        "mit Bereichen, Linien und Kopplern, aus Topologie und Gewerken die "
        "vollständige Gruppenadress-Struktur, und schliesslich die passenden "
        "Sensoren und deren Tastenbelegung. Wer ein bestehendes ETS6-Projekt hat, "
        "kann es stattdessen importieren, prüfen lassen und reorganisieren.")

    # ── 2 Zielgruppe ──────────────────────────────────────────────────────
    add_heading1(doc, "Für wen ist KNiX Arranger gemacht?")
    add_body(doc,
        "KNiX Arranger richtet sich an alle, die KNX-Anlagen planen, anbieten "
        "oder in Betrieb nehmen: KNX-Systemintegratoren und Elektroplanungsbüros, "
        "die ihre Projekte schneller und konsistenter ausarbeiten wollen; "
        "Gebäudeautomations-Ingenieure und Projektleiter, die den Ueberblick "
        "über grosse Adress- und Gerätestrukturen behalten müssen; sowie "
        "Neueinsteiger und Ausbildungsstätten, die anhand der automatisch "
        "erzeugten, normkonformen Strukturen die KNX-Planungsmethodik erlernen "
        "können. Allen gemeinsam ist, dass repetitive, fehleranfällige "
        "Routinearbeit an die Software delegiert wird, während die fachliche "
        "Entscheidung – welche Gewerke, welche Produkte, welche Funktionen – "
        "beim Planer bleibt.")

    # ── 3 Planungsprozess ────────────────────────────────────────────────
    add_heading1(doc, "Der Planungsprozess")
    add_body(doc,
        "Der Wizard führt durch die Gebäudestruktur (Gebäude, Flügel, "
        "Stockwerke, Zonen/Wohnungen und Räume – inklusive Unterstützung für "
        "Maisonette-Wohnungen über mehrere Stockwerke), die Elektroverteilungen "
        "(Haupt- und Unterverteilungen) und die Gewerke-Zuweisung pro Raum. "
        "Anschliessend konfiguriert eine Matrix-Ansicht die Tastereinheiten, "
        "bevor KNiX Arranger automatisch die KNX-Topologie berechnet: Bereiche, "
        "Linien, Linien- und Bereichskoppler sowie Speisegeräte werden nach den "
        "Regeln der KNX TP-Topologie und den KNX Swiss Projektrichtlinien "
        "vorgeschlagen und können bei Bedarf manuell angepasst werden.")
    add_body(doc,
        "Aus den Gewerken leitet das Programm die benötigten Aktoren ab – "
        "Schalt-, Dimm-, Jalousie- und Heizungsaktoren ebenso wie Gateways für "
        "DALI, Multimedia, Wärmepumpe oder E-Ladestationen – und ordnet sie den "
        "passenden Unterverteilungen zu. Vor der Gruppenadress-Generierung "
        "können Licht- und Jalousieszenen definiert werden, die automatisch in "
        "die Hauptgruppe 0 eingeplant werden. Danach generiert KNiX Arranger die "
        "vollständige, dreistufige Gruppenadress-Struktur (Haupt-, Mittel- und "
        "Untergruppen) inklusive Bezeichnungen und Datenpunkttypen – wahlweise "
        "in der kompakten Variante A (Rückmeldungen in derselben Mittelgruppe) "
        "oder der übersichtlicheren Variante B (separate Mittelgruppen für "
        "Rückmeldungen, typisch für Mehrfamilienhäuser).")
    add_body(doc,
        "Im letzten Planungsabschnitt ermittelt das Programm die erforderlichen "
        "Sensoren (Tastsensoren, Raumthermostate, Präsenzmelder, Systemsensoren "
        "wie Wetterstationen) und ordnet ihnen Funktionen zu. In der "
        "Bauherren-Beratung können diese Funktionen gemeinsam mit dem Bauherrn "
        "am Bildschirm besprochen und dessen Wünsche je Taste direkt erfasst "
        "werden – die Grundlage für das unterschriftsreife Bauherr-Formular. "
        "Am Ende steht der Export: als ETS6-kompatible CSV-Datei, als natives "
        "KNXPROJ-Projekt oder als vollständiges Revisionspaket.")

    # ── 4 Funktionsumfang ─────────────────────────────────────────────────
    add_heading1(doc, "Funktionsumfang im Ueberblick")

    add_heading2(doc, "Gebäudestruktur und Gewerke")
    add_body(doc,
        "Areale, Gebäude, Stockwerke, Zonen und Räume werden in einer "
        "übersichtlichen Baumstruktur erfasst; fertige Vorlagen für "
        "Einfamilienhaus, Mehrfamilienhaus oder Zweckbau beschleunigen den "
        "Einstieg. Ein Katalog mit über 40 Gewerke-Codes deckt Licht (inkl. "
        "Farbsteuerung und DALI), Jalousie, Heizung, Lüftung/Klima, Energie, "
        "Multimedia, Alarm und allgemeine Funktionen ab und kann pro Raum frei "
        "kombiniert werden.")

    add_heading2(doc, "Automatische Topologie- und Gruppenadress-Generierung")
    add_body(doc,
        "Aus der Anzahl der Aktoren und Sensoren berechnet KNiX Arranger "
        "automatisch eine normkonforme KNX-Topologie mit Bereichen, Linien, "
        "Kopplern und Speisegeräten. Darauf aufbauend entsteht die komplette "
        "Gruppenadress-Struktur inklusive Bezeichnungskonzept, Mittelgruppen-"
        "Zuordnung und Datenpunkttypen – ohne dass eine einzige Adresse von Hand "
        "vergeben werden muss. Manuell ergänzte Adressen bleiben bei einer "
        "Neuberechnung erhalten.")

    add_heading2(doc, "Aktor- und Sensor-Ermittlung")
    add_body(doc,
        "Auf Basis der Gewerke ermittelt das Programm die benötigten "
        "Aktortypen und Kanalzahlen je Unterverteilung sowie die passenden "
        "Sensoren samt Tastenanzahl. Bevorzugte Hersteller können hinterlegt "
        "werden, und zu jedem Gewerk, Aktor oder Sensor lassen sich "
        "Produktdatenblätter speichern, die später in die Dokumentation "
        "einfliessen.")

    add_heading2(doc, "Szenen, Zeitsteuerung, DALI und KNX Secure")
    add_body(doc,
        "Licht- und Jalousieszenen werden mit Vorlagen (Kino, Dinner, "
        "Abwesenheit usw.) angelegt und automatisch verdrahtet. Eine "
        "Zeitsteuerungs-Ansicht erlaubt Wochenprogramme mit Astrofunktionen "
        "(Sonnenauf- und -untergang, Dämmerung) und einem Feiertagskalender. "
        "Für Projekte mit DALI-Beleuchtung erkennt KNiX Arranger Gateways und "
        "Gruppen automatisch und unterstützt die Detailkonfiguration von "
        "Vorschaltgeräten und Szenen. Wo erhöhte Sicherheitsanforderungen "
        "bestehen, können KNX-Secure-Schlüssel für Linien und Geräte "
        "verwaltet werden.")

    add_heading2(doc, "Bauherren-Beratung")
    add_body(doc,
        "Eine eigene Ansicht erlaubt es, die Tastenbelegung direkt mit dem "
        "Bauherrn am Bildschirm durchzugehen: Für jede Taste kann ein "
        "Wunsch ausgewählt und je Raum eine Anmerkung erfasst werden. Diese "
        "Eingaben bleiben auch nach einer automatischen Neuberechnung des "
        "Projekts erhalten und fliessen direkt in das Bauherr-Formular ein.")

    add_heading2(doc, "Import, Validierung und Reorganisation")
    add_body(doc,
        "Bestehende ETS6-Projekte lassen sich als Gruppenadress-CSV, als "
        "Topologie-Report (XLSX) oder als natives KNXPROJ-Archiv importieren. "
        "Eine integrierte Validierung prüft das Projekt vor dem Export auf "
        "Fehler und Warnungen – von fehlenden Datenpunkttypen über "
        "Adresskonflikte bis zu Abweichungen von den KNX Swiss "
        "Projektrichtlinien – und schlägt Lösungen vor. Bestehende, "
        "gewachsene Gruppenadress-Strukturen können normgerecht reorganisiert "
        "werden.")

    add_heading2(doc, "Materialliste, Produktzuweisung und Kundenofferten")
    add_body(doc,
        "Alle ermittelten Geräte laufen in einer Materialliste zusammen, in "
        "der konkrete Produkte zugewiesen werden – einzeln oder per "
        "Typ-Batch-Zuweisung aus einem lokalen Produktkatalog. Auf dieser "
        "Grundlage erstellt KNiX Arranger automatisch Kundenofferten mit den "
        "im Firmenprofil hinterlegten Stundensätzen für Montage, "
        "Programmierung und Inbetriebnahme.")

    # ── 5 Zeitersparnis ───────────────────────────────────────────────────
    add_heading1(doc, "Wo KNiX Arranger viel Zeit einspart")
    add_body(doc,
        "Der grösste Zeitgewinn entsteht dort, wo klassischerweise repetitive "
        "Handarbeit anfällt: Statt hunderte bis tausende Gruppenadressen "
        "einzeln zu benennen, zu nummerieren und mit Datenpunkttypen zu "
        "versehen, generiert KNiX Arranger die komplette Struktur in Sekunden "
        "– konsistent nach demselben Bezeichnungskonzept für das gesamte "
        "Projekt. Auch die Topologieplanung, die normalerweise eine sorgfältige "
        "manuelle Abstimmung von Teilnehmerzahlen, Linien- und "
        "Bereichsgrenzen erfordert, erfolgt automatisch und wird bei "
        "Aenderungen am Gebäude oder den Gewerken neu berechnet.")
    add_body(doc,
        "Auch die Sensor- und Aktor-Ermittlung, die sonst raumweise von Hand "
        "durchgegangen werden muss, geschieht automatisch aus den zugewiesenen "
        "Gewerken – inklusive Tastenanzahl und Funktionszuordnung per "
        "Auto-Assign. Wo früher mehrere Excel-Tabellen für Materialliste, "
        "Bauherr-Formular und Offerte parallel gepflegt wurden, entstehen diese "
        "Dokumente jetzt direkt und konsistent aus denselben Projektdaten. "
        "Aenderungen – etwa ein zusätzlicher Raum oder ein weiteres Gewerk – "
        "wirken sich automatisch auf Topologie, Gruppenadressen, "
        "Geräteliste und Dokumentation aus, ohne dass jede Folge manuell "
        "nachgezogen werden muss.")

    # ── 6 Dokumentation ───────────────────────────────────────────────────
    add_heading1(doc, "Dokumentation auf Knopfdruck")
    add_body(doc,
        "Am Ende der Planung steht nicht nur die Projektdatei, sondern ein "
        "vollständiger Satz an Unterlagen für Bauherr, Elektroplaner und die "
        "eigene Ablage. KNiX Arranger kann folgende Dokumente erzeugen:")
    add_bullet(doc,
        "CSV-Export: ETS6-kompatibler Gruppenadress-Export für den direkten "
        "Import in die ETS6 (Extras -> Gruppenadressbericht -> Import).")
    add_bullet(doc,
        "KNXPROJ-Export: natives ETS6-Projektformat mit Topologie und "
        "Gruppenadressen für den vollständigen Projektimport in einem "
        "Schritt.")
    add_bullet(doc,
        "Revisionspaket (ZIP): GA-Uebersicht als Excel-Datei, Topologie-Plan, "
        "Belegungsplan, Kabellängen-Report und Produktdatenblätter in einem "
        "Paket.")
    add_bullet(doc,
        "Bauherr-Formular (XLSX): die vollständige Tastenbelegung pro Raum "
        "zur Prüfung und Unterschrift durch den Bauherrn.")
    add_bullet(doc,
        "Abnahmeprotokoll: Vordruck für die Projektübergabe und "
        "Inbetriebnahme.")
    add_bullet(doc,
        "Kundenofferte: automatisch kalkuliertes Angebotsdokument auf Basis "
        "der Materialliste und der hinterlegten Stundensätze.")
    add_bullet(doc,
        "Validierungsbericht: Liste aller Fehler, Warnungen und Hinweise mit "
        "betroffener Gruppenadresse und Lösungsvorschlag.")
    add_body(doc,
        "Alle Dokumente werden mit dem im Firmenprofil hinterlegten "
        "Briefkopf, Logo und Kontaktdaten erstellt, sodass sie ohne weitere "
        "Nachbearbeitung an Kunden und Partner weitergegeben werden können.")

    # ── 7 Fazit ───────────────────────────────────────────────────────────
    add_heading1(doc, "Fazit")
    add_body(doc,
        "KNiX Arranger nimmt KNX-Systemintegratoren die zeitaufwendige, "
        "fehleranfällige Routinearbeit der Projektplanung ab und sorgt dabei "
        "für durchgängig normkonforme, gut dokumentierte Projekte – von der "
        "ersten Gebäudestruktur bis zur unterschriftsreifen Bauherrenunterlage "
        "und dem ETS6-Import. Das gewonnene Zeitbudget steht für das, worauf es "
        "wirklich ankommt: die fachliche Beratung und die Qualität der "
        "Installation vor Ort.")

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Müller SmartHome & EnergieManagement")
    run.font.size = Pt(11)
    run.font.color.rgb = C_GREEN_DARK

    return doc


if __name__ == "__main__":
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

    doc = build()
    doc.save(OUT_PATH)
    print(f"Gespeichert: {OUT_PATH}")
    print(f"Absätze: {len(doc.paragraphs)}")
