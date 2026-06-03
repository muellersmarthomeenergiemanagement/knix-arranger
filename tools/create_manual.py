"""
Erstellt die Beta-Tester-Anleitung als Word-Dokument (.docx)

Ausfuehren:
    python tools/create_manual.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_PATH = Path(__file__).parent.parent / "KNiX_Arranger_Betatester_Anleitung.docx"

DARK_GREEN = RGBColor(0x1B, 0x5E, 0x20)
MID_GREEN  = RGBColor(0x2E, 0x7D, 0x32)
GREY       = RGBColor(0x55, 0x55, 0x55)


def set_heading_color(paragraph, color):
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_colored_heading(doc, text, level, color=DARK_GREEN):
    p = doc.add_heading(text, level=level)
    set_heading_color(p, color)
    return p


def add_info_box(doc, text):
    """Fügt einen grauen Hinweiskasten ein."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "E8F5E9")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    return p


def add_step(doc, number, title, description):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.color.rgb = MID_GREEN
    doc.add_paragraph(description, style="Normal")


def build():
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Schriftart global
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── Titelseite ─────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KNiX Arranger")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_GREEN

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Betatester-Anleitung  ·  Version 1.0.0")
    run.font.size = Pt(14)
    run.font.color.rgb = GREY

    doc.add_paragraph()

    add_info_box(doc,
        "Vielen Dank, dass Sie als Betatester an der Entwicklung von KNiX Arranger mitwirken!\n"
        "Diese Anleitung führt Sie durch die Installation, Lizenzaktivierung und die wichtigsten "
        "Funktionen des Programms. Ihr Feedback hilft uns, das Produkt zu verbessern."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mueller SmartHome & EnergieManagement")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("info@muellersmarthomeenergiemanagement.ch")
    run2.font.size = Pt(11)
    run2.font.color.rgb = MID_GREEN

    doc.add_page_break()

    # ── 1. Systemvoraussetzungen ───────────────────────────────────────────────
    add_colored_heading(doc, "1  Systemvoraussetzungen", 1)

    t = doc.add_table(rows=5, cols=2)
    t.style = "Table Grid"
    data = [
        ("Betriebssystem", "Windows 10 oder Windows 11 (64-Bit)"),
        ("Prozessor",      "Intel/AMD, 1.5 GHz oder schneller"),
        ("Arbeitsspeicher","4 GB RAM (8 GB empfohlen)"),
        ("Festplatte",     "ca. 300 MB freier Speicher"),
        ("Bildschirm",     "Mindestauflösung 1280 × 768"),
    ]
    for i, (label, value) in enumerate(data):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── 2. Installation ────────────────────────────────────────────────────────
    add_colored_heading(doc, "2  Installation", 1)

    doc.add_paragraph(
        "Sie haben per E-Mail zwei Dateien erhalten:"
    )
    ul = doc.add_paragraph(style="List Bullet")
    ul.add_run("KNiX_Arranger_vX.X.X.zip").bold = True
    ul.add_run("  –  das Programm als ZIP-Archiv")

    ul2 = doc.add_paragraph(style="List Bullet")
    ul2.add_run("IhrName_JJJJ-MM-TT.knxlic").bold = True
    ul2.add_run("  –  Ihre persönliche Lizenzdatei")

    doc.add_paragraph()
    add_colored_heading(doc, "2.1  ZIP entpacken", 2)
    doc.add_paragraph(
        "Klicken Sie mit der rechten Maustaste auf die ZIP-Datei und wählen Sie "
        "«Alle extrahieren…». Wählen Sie einen Zielordner, z. B. "
        "C:\\Programme\\KNiX Arranger, und klicken Sie auf «Extrahieren»."
    )

    add_colored_heading(doc, "2.2  Programm starten", 2)
    doc.add_paragraph(
        "Öffnen Sie den entpackten Ordner und doppelklicken Sie auf "
        "KNX_Arranger.exe. "
        "Optional können Sie eine Verknüpfung auf dem Desktop anlegen "
        "(Rechtsklick → Senden an → Desktop)."
    )

    add_info_box(doc,
        "Hinweis: Wenn Windows SmartScreen erscheint («Unbekannter Herausgeber»), "
        "klicken Sie auf «Weitere Informationen» und dann «Trotzdem ausführen». "
        "Dies ist beim ersten Start eines nicht signierten Programms normal."
    )

    doc.add_paragraph()

    # ── 3. Lizenzaktivierung ───────────────────────────────────────────────────
    add_colored_heading(doc, "3  Lizenzaktivierung", 1)

    doc.add_paragraph(
        "Beim ersten Start erscheint der Lizenz-Dialog. Gehen Sie wie folgt vor:"
    )

    for num, (title, desc) in enumerate([
        ("EULA akzeptieren",
         "Lesen Sie die Endnutzer-Lizenzvereinbarung und klicken Sie auf «Akzeptieren»."),
        ("Lizenzdatei auswählen",
         "Im Lizenz-Dialog klicken Sie auf «Lizenzdatei auswählen…» und wählen die "
         "mitgelieferte .knxlic-Datei aus Ihrem Download-Ordner aus."),
        ("Bestätigung",
         "Nach erfolgreichem Import erscheint die Meldung «Lizenz erfolgreich importiert». "
         "Die Anwendung startet nun vollständig."),
    ], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{title}:  ").bold = True
        p.add_run(desc)

    add_info_box(doc,
        "Testlizenz: Ihre Betatester-Lizenz ist 30 Tage gültig. "
        "10 Tage vor Ablauf erscheint ein Hinweis beim Programmstart. "
        "Melden Sie sich bei uns, falls Sie eine Verlängerung benötigen."
    )

    doc.add_paragraph()

    # ── 4. Programmübersicht ───────────────────────────────────────────────────
    add_colored_heading(doc, "4  Programmübersicht", 1)

    doc.add_paragraph(
        "Das Hauptfenster von KNiX Arranger besteht aus drei Bereichen:"
    )

    t2 = doc.add_table(rows=3, cols=2)
    t2.style = "Table Grid"
    areas = [
        ("Seitenleiste (links)",
         "Navigation zwischen den verschiedenen Ansichten: Übersicht, Gebäude, "
         "Topologie, Gruppenadressen, Materialliste, Szenen, Zeitsteuerung und mehr."),
        ("Hauptbereich (Mitte)",
         "Zeigt die jeweils gewählte Ansicht. Hier finden statt: Dateneingabe, "
         "Tabellen, Baumstrukturen, Diagramme."),
        ("Statusleiste (unten)",
         "Zeigt Projektname, Anzahl Gruppenadressen und Validierungsstatus."),
    ]
    for i, (label, value) in enumerate(areas):
        t2.rows[i].cells[0].text = label
        t2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t2.rows[i].cells[1].text = value

    doc.add_paragraph()
    add_colored_heading(doc, "4.1  Menüleiste", 2)

    t3 = doc.add_table(rows=4, cols=2)
    t3.style = "Table Grid"
    menus = [
        ("Datei",      "Neues Projekt, Öffnen, Speichern, Import (ETS6/CSV/KNXPROJ), Export"),
        ("Bearbeiten", "Rückgängig (Ctrl+Z), Wiederholen (Ctrl+Y), Einstellungen"),
        ("Ansicht",    "Wizard starten (Ctrl+W), Projekt validieren (Ctrl+V)"),
        ("Hilfe",      "Hilfe (F1), Tour, Benutzerhandbuch, Updates, Lizenz"),
    ]
    for i, (label, value) in enumerate(menus):
        t3.rows[i].cells[0].text = label
        t3.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t3.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── 5. Neues Projekt anlegen ───────────────────────────────────────────────
    add_colored_heading(doc, "5  Neues Projekt anlegen", 1)

    doc.add_paragraph(
        "Beim Start erscheint der Willkommens-Dialog. Wählen Sie «Neues Projekt» "
        "oder klicken Sie im Menü auf Datei → Neues Projekt (Ctrl+N)."
    )
    doc.add_paragraph(
        "Geben Sie Projektname, Auftraggeber und Projektstandort ein und bestätigen "
        "Sie mit «Erstellen». Das Projekt wird als .knxarr-Datei gespeichert."
    )

    doc.add_paragraph()

    # ── 6. Der Planungs-Wizard ─────────────────────────────────────────────────
    add_colored_heading(doc, "6  Der Planungs-Wizard", 1)

    doc.add_paragraph(
        "Der Wizard führt Sie in 13 Schritten durch die vollständige KNX-Projektplanung. "
        "Starten Sie ihn über Ansicht → Wizard starten (Ctrl+W) oder über die Übersichtsseite."
    )

    steps = [
        ("Schritt 1 – Gebäudestruktur",
         "Definieren Sie Gebäude, Flügel und Stockwerke. Wählen Sie ob es sich um "
         "ein Einfamilienhaus (EFH) oder Mehrfamilienhaus (MFH) handelt."),
        ("Schritt 2 – Wohnungen / Zonen",
         "Legen Sie Wohnungen (MFH) oder Zonen (EFH/Büro) an und ordnen Sie ihnen "
         "Stockwerke zu. Eine Zone kann mehrere Stockwerke umfassen (Maisonette)."),
        ("Schritt 3 – Räume",
         "Fügen Sie Räume pro Zone und Stockwerk hinzu. Die Raumnummerierung "
         "erfolgt automatisch nach dem gewählten Schema."),
        ("Schritt 4 – Elektroverteilungen",
         "Erfassen Sie Haupt- und Unterverteilungen (HV/UV) und ordnen Sie diese "
         "den Gebäudeteilen zu."),
        ("Schritt 5 – Gewerke",
         "Weisen Sie jedem Raum die gewünschten KNX-Gewerke zu (z. B. Licht schalten, "
         "Jalousie, Heizung, DALI). Wählen Sie Varianten und Besonderheiten pro Raum."),
        ("Schritt 6 – Tastereinheiten",
         "Konfigurieren Sie die Tastereinheiten pro Raum in der Matrix-Ansicht. "
         "Jede Zeile entspricht einem Gewerk, jede Spalte einer Tastereinheit."),
        ("Schritt 7 – Topologie",
         "KNiX Arranger berechnet die KNX-Topologie automatisch. Prüfen Sie "
         "Bereiche, Linien und empfohlene Linienkoppler."),
        ("Schritt 8 – Aktoren",
         "Übersicht der ermittelten Aktoren nach Gerätetyp, Kanal und Raum. "
         "Kontrollieren Sie die automatische Kanalzuweisung."),
        ("Schritt 9 – Szenen",
         "Definieren Sie Licht- und Jalousie-Szenen vor der Gruppenadress-Generierung."),
        ("Schritt 10 – Gruppenadressen",
         "Die Gruppenadressen werden automatisch generiert. Prüfen und ergänzen Sie "
         "bei Bedarf einzelne Adressen."),
        ("Schritt 11 – Sensoren",
         "Übersicht der ermittelten Sensoren, Taster und Systemsensoren "
         "mit ihren Funktionszuordnungen."),
        ("Schritt 12 – Funktionsdefinition",
         "Das Bauherr-Formular zeigt alle Bedienelemente und deren zugeordneten "
         "Funktionen – ideal als Grundlage für die Inbetriebnahme."),
        ("Schritt 13 – Export",
         "Exportieren Sie das Projekt als CSV für ETS6, als natives .knxproj-Format "
         "oder als Revisionspaket (PDF, Excel)."),
    ]

    for num, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{title}").bold = True
        doc.add_paragraph(desc)

    doc.add_paragraph()

    # ── 7. Wichtige Ansichten ──────────────────────────────────────────────────
    add_colored_heading(doc, "7  Wichtige Ansichten in der Seitenleiste", 1)

    views = [
        ("Übersicht",          "Projektinformationen, Fortschrittsanzeige, Schnellzugriff auf Wizard"),
        ("Gebäude",            "Baumstruktur: Gebäude → Flügel → Stockwerk → Zone → Raum"),
        ("Topologie",          "KNX-Linienkoppler, Speisegeräte, Bereichskoppler mit Leistungsberechnung"),
        ("Gruppenadressen",    "Tabellarische Übersicht aller GAs, filterbar nach Gewerk"),
        ("Verknüpfungsmatrix", "Welcher Sensor sendet auf welche GA? Übersicht aller CO-Verknüpfungen"),
        ("Materialliste",      "Ermittelte Geräte mit Typ, Kanal, Raum – Produktzuweisung möglich"),
        ("Szenen",             "Szenen-Verwaltung mit Aktorwerten"),
        ("Zeitsteuerung",      "Wochenprogramme, Astrofunktionen, Feiertagskalender"),
        ("DALI-Konfiguration", "DALI-Gruppen und EVGs, automatisch aus importierten GAs ermittelt"),
        ("Validierung",        "Prüft das Projekt auf Fehler und Warnungen vor dem Export"),
        ("Export",             "Revisionspaket, CSV, KNXPROJ, Belegungsplan, Kabelwege"),
    ]

    t4 = doc.add_table(rows=len(views), cols=2)
    t4.style = "Table Grid"
    for i, (label, value) in enumerate(views):
        t4.rows[i].cells[0].text = label
        t4.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t4.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── 8. Tastenkürzel ───────────────────────────────────────────────────────
    add_colored_heading(doc, "8  Wichtige Tastenkürzel", 1)

    shortcuts = [
        ("Ctrl+N",       "Neues Projekt"),
        ("Ctrl+O",       "Projekt öffnen"),
        ("Ctrl+S",       "Speichern"),
        ("Ctrl+Shift+S", "Speichern unter"),
        ("Ctrl+W",       "Wizard starten"),
        ("Ctrl+I",       "Import (CSV / XLSX / KNXPROJ)"),
        ("Ctrl+E",       "CSV-Export"),
        ("Ctrl+Z",       "Rückgängig"),
        ("Ctrl+Y",       "Wiederholen"),
        ("Ctrl+V",       "Validieren"),
        ("F1",           "Hilfe anzeigen"),
    ]

    t5 = doc.add_table(rows=len(shortcuts), cols=2)
    t5.style = "Table Grid"
    for i, (key, desc) in enumerate(shortcuts):
        t5.rows[i].cells[0].text = key
        t5.rows[i].cells[0].paragraphs[0].runs[0].font.name = "Courier New"
        t5.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t5.rows[i].cells[1].text = desc

    doc.add_paragraph()

    # ── 9. Bekannte Einschränkungen (Beta) ────────────────────────────────────
    add_colored_heading(doc, "9  Bekannte Einschränkungen (Beta-Version)", 1)

    doc.add_paragraph(
        "Als Betatester erhalten Sie eine Vorabversion. Folgende Punkte sind bekannt:"
    )
    for item in [
        "Einzelne Exportfunktionen können noch unvollständig sein.",
        "Die integrierte Hilfedokumentation wird noch ausgebaut.",
        "Das Programm wurde noch nicht auf allen Systemkonfigurationen getestet.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # ── 10. Feedback ──────────────────────────────────────────────────────────
    add_colored_heading(doc, "10  Feedback erwünscht!", 1)

    doc.add_paragraph(
        "Ihr Feedback ist für uns sehr wertvoll. Bitte melden Sie uns:"
    )
    for item in [
        "Fehler und unerwartetes Verhalten (mit kurzer Beschreibung der Schritte)",
        "Unklare oder fehlende Funktionen",
        "Verbesserungsvorschläge zur Bedienbarkeit",
        "Allgemeine Eindrücke zur Anwendung",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Feedback per E-Mail an:  ").bold = True
    run = p.add_run("info@muellersmarthomeenergiemanagement.ch")
    run.font.color.rgb = MID_GREEN
    run.bold = True

    doc.add_paragraph(
        "Bitte geben Sie in der Betreffzeile «KNiX Arranger Beta» an. "
        "Für Absturz-Reports finden Sie unter %APPDATA%\\KNiX Arranger\\ "
        "eine Log-Datei, die Sie uns beilegen können."
    )

    doc.add_paragraph()
    add_info_box(doc,
        "Herzlichen Dank für Ihre Unterstützung!\n"
        "Mueller SmartHome & EnergieManagement"
    )

    doc.save(OUT_PATH)
    print(f"Anleitung gespeichert: {OUT_PATH}")


if __name__ == "__main__":
    build()
