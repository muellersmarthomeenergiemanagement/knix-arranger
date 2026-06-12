"""
Erstellt die verbesserte Betatester-Anleitung (KNiX_Arranger_Betatester_Anleitung.docx).

Verbesserungen gegenüber v1.0.0:
- Wizard-Schritttitel korrigiert (Schritt 6 + 11)
- Alle 13 Wizard-Schritte mit ausführlichen Beschreibungen
- Lizenzdauer: Ablaufdatum statt "30 Tage"
- Neues Kapitel: ETS6-Import-Workflow
- Neues Kapitel: Hilfe-System (F1, ?-Button, Tour)
- Neues Kapitel: Troubleshooting
- Kapitel Seitenleiste: Wann-/Warum-Spalte ergänzt
- Feedback: strukturiertes Template hinzugefügt
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Farben ────────────────────────────────────────────────────────────────────
C_GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)
C_GREEN       = RGBColor(0x2E, 0x7D, 0x32)
C_BLUE        = RGBColor(0x36, 0x5F, 0x91)
C_GREY        = RGBColor(0x55, 0x55, 0x55)
C_ORANGE      = RGBColor(0xE6, 0x51, 0x00)
C_BLACK       = RGBColor(0x00, 0x00, 0x00)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def add_heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)


def add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)


def add_body(doc: Document, text: str, bold_parts: list[str] | None = None) -> None:
    """Fügt Normal-Absatz hinzu. bold_parts = Texte die fett sein sollen."""
    p = doc.add_paragraph(style="Normal")
    if not bold_parts:
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        # Einfaches Splitting für fette Teile
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx >= 0:
                before = remaining[:idx]
                after  = remaining[idx + len(bp):]
                if before:
                    r = p.add_run(before)
                    r.font.size = Pt(10)
                r = p.add_run(bp)
                r.bold = True
                r.font.size = Pt(10)
                remaining = after
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(10)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(10)


def add_tip(doc: Document, text: str) -> None:
    """Tipp-Box: grüner Einzug-Absatz."""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run("Tipp: ")
    run.bold = True
    run.font.color.rgb = C_GREEN
    run.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.color.rgb = C_GREEN
    r2.font.size = Pt(10)


def add_note(doc: Document, text: str) -> None:
    """Hinweis-Box: orange Einzug."""
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run("Hinweis: ")
    run.bold = True
    run.font.color.rgb = C_ORANGE
    run.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.color.rgb = C_ORANGE
    r2.font.size = Pt(10)


def add_table_2col(doc: Document, rows: list[tuple[str, str]],
                   header: tuple[str, str] | None = None) -> None:
    """2-spaltige Tabelle mit optionalem Header."""
    table = doc.add_table(rows=0, cols=2, style="Table Grid")
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)

    if header:
        row = table.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            # Hintergrundfarbe Header
            set_cell_bg(cell, "365F91")
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (col1, col2) in enumerate(rows):
        row = table.add_row()
        c1, c2 = row.cells[0], row.cells[1]
        c1.text = col1
        c2.text = col2
        for c in (c1, c2):
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_bg(c1, "EEF2F7")
            set_cell_bg(c2, "EEF2F7")


def set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_wizard_step(doc: Document, number: int, title: str, description: str,
                    eingabe: list[str] | None = None,
                    berechnet: list[str] | None = None,
                    tipp: str | None = None) -> None:
    """Formatiert einen Wizard-Schritt einheitlich."""
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"Schritt {number} – {title}")
    run.bold = True
    run.font.size = Pt(10)

    add_body(doc, description)

    if eingabe:
        p2 = doc.add_paragraph(style="Normal")
        p2.paragraph_format.left_indent = Cm(0.5)
        r = p2.add_run("Was Sie eingeben:")
        r.bold = True
        r.font.size = Pt(10)
        for item in eingabe:
            pi = doc.add_paragraph(style="List Bullet")
            pi.paragraph_format.left_indent = Cm(0.5)
            pi.add_run(item).font.size = Pt(10)

    if berechnet:
        p3 = doc.add_paragraph(style="Normal")
        p3.paragraph_format.left_indent = Cm(0.5)
        r = p3.add_run("Was automatisch berechnet wird:")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_GREEN
        for item in berechnet:
            pi = doc.add_paragraph(style="List Bullet")
            pi.paragraph_format.left_indent = Cm(0.5)
            pi.add_run(item).font.size = Pt(10)

    if tipp:
        add_tip(doc, tipp)

    doc.add_paragraph()  # Leerzeile nach Schritt


# ── Dokument aufbauen ─────────────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # Seiteneinstellungen übernehmen
    section = doc.sections[0]
    section.page_width  = Cm(21.6)
    section.page_height = Cm(27.9)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Titelseite ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KNiX Arranger")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = C_GREEN_DARK

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Betatester-Anleitung  ·  Version 1.1.2")
    run.font.size = Pt(14)
    run.font.color.rgb = C_GREY

    doc.add_paragraph()

    p = doc.add_paragraph(style="Normal")
    run = p.add_run(
        "Vielen Dank, dass Sie als Betatester an der Entwicklung von KNiX Arranger mitwirken! "
        "Diese Anleitung führt Sie durch die Installation, Lizenzaktivierung und die wichtigsten "
        "Funktionen des Programms. Ihr Feedback hilft uns, das Produkt zu verbessern."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = C_GREY

    doc.add_paragraph()

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Mueller SmartHome & EnergieManagement").font.size = Pt(11)
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("info@muellersmarthomeenergiemanagement.ch")
    run.font.size = Pt(11)
    run.font.color.rgb = C_GREEN

    doc.add_paragraph()

    # ── 1  Systemvoraussetzungen ──────────────────────────────────────────────
    add_heading1(doc, "1  Systemvoraussetzungen")

    add_table_2col(doc, [
        ("Betriebssystem",  "Windows 10 oder Windows 11 (64-Bit)"),
        ("Prozessor",       "Intel/AMD, 1.5 GHz oder schneller"),
        ("Arbeitsspeicher", "4 GB RAM (8 GB empfohlen)"),
        ("Festplatte",      "ca. 300 MB freier Speicher"),
        ("Bildschirm",      "Mindestauflösung 1280 × 768 (empfohlen: 1920 × 1080)"),
        ("Netzwerk",        "Keine Internetverbindung erforderlich (optional für Updates)"),
    ])
    doc.add_paragraph()

    # ── 2  Installation ───────────────────────────────────────────────────────
    add_heading1(doc, "2  Installation")
    add_body(doc, "Das Programm steht zum Download bereit unter:")
    add_bullet(doc, "https://github.com/muellersmarthomeenergiemanagement/knix-arranger-releases/releases/latest")
    add_bullet(doc, "IhrName_JJJJ-MM-TT.knxlic  –  Ihre persönliche Lizenzdatei")
    doc.add_paragraph()

    add_heading2(doc, "2.1  ZIP entpacken")
    add_body(doc,
        "Klicken Sie mit der rechten Maustaste auf die ZIP-Datei und wählen Sie "
        "«Alle extrahieren…». Wählen Sie einen Zielordner, z. B. "
        "C:\\Programme\\KNiX Arranger, und klicken Sie auf «Extrahieren».")
    doc.add_paragraph()

    add_heading2(doc, "2.2  Programm starten")
    add_body(doc,
        "Öffnen Sie den entpackten Ordner und doppelklicken Sie auf KNX_Arranger.exe. "
        "Optional können Sie eine Verknüpfung auf dem Desktop anlegen "
        "(Rechtsklick → Senden an → Desktop).")
    add_note(doc,
        "Wenn Windows SmartScreen erscheint («Unbekannter Herausgeber»), klicken Sie auf "
        "«Weitere Informationen» und dann «Trotzdem ausführen». Dies ist beim ersten Start "
        "eines nicht signierten Programms normal und kein Sicherheitsrisiko.")
    doc.add_paragraph()

    # ── 3  Lizenzaktivierung ──────────────────────────────────────────────────
    add_heading1(doc, "3  Lizenzaktivierung")
    add_body(doc, "Beim ersten Start erscheint der Lizenz-Dialog. Gehen Sie wie folgt vor:")
    add_numbered(doc,
        "EULA akzeptieren:  Lesen Sie die Endnutzer-Lizenzvereinbarung und klicken "
        "Sie auf «Akzeptieren».")
    add_numbered(doc,
        "Lizenzdatei auswählen:  Im Lizenz-Dialog klicken Sie auf «Lizenzdatei "
        "auswählen…» und wählen die mitgelieferte .knxlic-Datei aus Ihrem "
        "Download-Ordner aus.")
    add_numbered(doc,
        "Bestätigung:  Nach erfolgreichem Import erscheint die Meldung «Lizenz "
        "erfolgreich importiert». Die Anwendung startet nun vollständig.")
    doc.add_paragraph()
    add_body(doc,
        "Betatester-Lizenz: Ihre Lizenz ist gültig bis 1. Juli 2026. "
        "14 Tage vor Ablauf erscheint beim Programmstart ein Hinweis. "
        "Melden Sie sich bei uns, falls Sie eine Verlängerung benötigen.")
    add_note(doc,
        "Die Lizenzdatei ist personalisiert und nicht übertragbar. "
        "Bitte geben Sie sie nicht an Dritte weiter.")
    doc.add_paragraph()

    # ── 3.1  Ersteinrichtung: Firmendaten ─────────────────────────────────────
    add_heading2(doc, "3.1  Ersteinrichtung: Firmendaten eingeben")

    # Hervorgehobene Box – bewusst auffällig gestaltet
    p_box = doc.add_paragraph(style="Normal")
    p_box.paragraph_format.left_indent  = Cm(0.6)
    p_box.paragraph_format.right_indent = Cm(0.6)
    p_box.paragraph_format.space_before = Pt(6)
    p_box.paragraph_format.space_after  = Pt(6)
    run_icon = p_box.add_run("⚠  ")
    run_icon.bold = True
    run_icon.font.size = Pt(11)
    run_icon.font.color.rgb = C_ORANGE
    run_bold = p_box.add_run(
        "Wichtig: Geben Sie direkt nach der ersten Lizenzaktivierung Ihre Firmendaten ein!"
    )
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    run_bold.font.color.rgb = C_ORANGE

    add_body(doc,
        "Die Firmendaten erscheinen auf allen exportierten Dokumenten: "
        "Revisionspaket, Bauherr-Formular, Abnahmeprotokoll, Topologie-Plan und Offerten. "
        "Ohne ausgefülltes Firmenprofil erscheinen auf den Dokumenten nur Platzhalter.")
    doc.add_paragraph()

    add_body(doc,
        "Öffnen Sie die Einstellungen über:  Bearbeiten → Einstellungen → Tab «Firmenprofil»")
    doc.add_paragraph()

    add_table_2col(doc,
        header=("Feld", "Was eingeben"),
        rows=[
            ("Firmenname",     "Vollständiger Firmenname (erscheint als Kopfzeile auf allen Dokumenten)"),
            ("Adresse",        "Strasse, PLZ, Ort"),
            ("Telefon",        "Direkte Nummer des zuständigen Projektleiters"),
            ("E-Mail",         "Kontakt-E-Mail (erscheint im Revisionspaket)"),
            ("Website",        "Optional – z. B. www.musterfirma.ch"),
            ("Bearbeiter Name","Ihr Name als zuständiger KNX-Systemintegrator"),
            ("Funktion",       "Ihre Rolle, z. B. «KNX-Systemintegrator»"),
            ("Firmenlogo",     "PNG oder JPG, wird auf Deckblättern eingebunden. "
                               "Empfohlene Grösse: 300 × 150 px, weisser Hintergrund."),
        ]
    )
    doc.add_paragraph()

    add_body(doc,
        "Im Tab «Stundensätze» können Sie ausserdem die Ansätze für "
        "Montage, Programmierung und Inbetriebnahme hinterlegen – "
        "diese werden für die automatische Kundenofferten-Kalkulation verwendet.")
    add_tip(doc,
        "Die Firmendaten werden einmalig gespeichert und gelten für alle Projekte. "
        "Sie müssen sie nicht pro Projekt neu eingeben.")
    doc.add_paragraph()

    # ── 4  Programmübersicht ──────────────────────────────────────────────────
    add_heading1(doc, "4  Programmübersicht")
    add_body(doc, "Das Hauptfenster von KNiX Arranger besteht aus drei Bereichen:")

    add_table_2col(doc, [
        ("Seitenleiste (links)",
         "Navigation zwischen den verschiedenen Ansichten: Übersicht, Gebäude, Topologie, "
         "Gruppenadressen, Materialliste, Szenen, Zeitsteuerung und mehr."),
        ("Hauptbereich (Mitte)",
         "Zeigt die jeweils gewählte Ansicht. Hier finden statt: Dateneingabe, Tabellen, "
         "Baumstrukturen, Diagramme."),
        ("Statusleiste (unten)",
         "Zeigt Projektname, Anzahl Gruppenadressen und Validierungsstatus."),
    ])
    doc.add_paragraph()

    add_heading2(doc, "4.1  Menüleiste")
    add_table_2col(doc, [
        ("Datei",    "Neues Projekt, Öffnen, Speichern, Import (ETS6/CSV/KNXPROJ), Export"),
        ("Bearbeiten", "Rückgängig (Ctrl+Z), Wiederholen (Ctrl+Y), Einstellungen"),
        ("Ansicht",  "Wizard starten (Ctrl+W), Projekt validieren (Ctrl+V)"),
        ("Hilfe",    "Hilfe (F1), Tour, Benutzerhandbuch, Updates, Lizenz"),
    ])
    doc.add_paragraph()

    add_heading2(doc, "4.2  Auto-Speichern")
    add_body(doc,
        "Das Programm speichert das Projekt automatisch 30 Sekunden nach der letzten Änderung. "
        "Die Datei wird im selben Ordner gespeichert, in dem das Projekt geöffnet wurde. "
        "Das Projektformat ist .knxarr (proprietäres JSON-Format).")
    add_tip(doc,
        "Verwenden Sie Ctrl+S um jederzeit manuell zu speichern – "
        "besonders vor einem Export oder grösseren Änderungen.")
    doc.add_paragraph()

    # ── 5  Neues Projekt anlegen ──────────────────────────────────────────────
    add_heading1(doc, "5  Neues Projekt anlegen")
    add_body(doc,
        "Beim allerersten Programmstart fragt KNiX Arranger einmalig nach einem "
        "Arbeitsverzeichnis (Workspace) – dem Ordner, in dem künftig alle Projekte "
        "abgelegt werden. Diesen Pfad können Sie jederzeit über "
        "Bearbeiten → Einstellungen → Tab «Arbeitsverzeichnis» ändern.")
    add_body(doc,
        "Für ein neues Projekt wählen Sie im Willkommens-Dialog «Neues Projekt» oder "
        "klicken im Menü auf Datei → Neues Projekt (Ctrl+N). Im Dialog geben Sie an:")
    add_bullet(doc, "Projektname (z. B. «Neubau EFH Mueller»)")
    add_bullet(doc, "Projektnummer (z. B. «2024-001»)")
    add_bullet(doc, "Gebäudetyp / Vorlage: Leeres Projekt, Einfamilienhaus (EFH), "
                    "Mehrfamilienhaus (MFH) oder Zweckbau")
    add_bullet(doc, "Mittelgruppen-Variante: A (Rückmeldungen in gleicher MG, Standard) "
                    "oder B (Rückmeldungen in separater MG 6/7)")
    add_body(doc,
        "Der Dialog zeigt direkt den künftigen Projektordner an. Mit «Projekt erstellen» "
        "wird folgende Ordnerstruktur im Arbeitsverzeichnis angelegt:")
    add_bullet(doc, "{Projektname}\\{Projektname}.knxarr  –  die Projektdatei")
    add_bullet(doc, "{Projektname}\\Revisionen\\  –  für exportierte Revisionspakete")
    add_bullet(doc, "{Projektname}\\Berichte\\  –  für weitere Berichte "
                    "(Offerten, Abnahmeprotokolle usw.)")
    add_tip(doc,
        "Nach dem Erstellen erscheint automatisch der Onboarding-Wizard, der Sie durch "
        "die ersten Schritte führt. Sie können ihn jederzeit über Hilfe → Erste Schritte "
        "erneut aufrufen.")
    doc.add_paragraph()

    # ── 6  ETS6-Import ────────────────────────────────────────────────────────
    add_heading1(doc, "6  Bestehendes ETS6-Projekt importieren")
    add_body(doc,
        "Wenn Sie bereits ein bestehendes KNX-Projekt in ETS6 haben, können Sie es in "
        "den KNiX Arranger importieren. Gehen Sie dazu auf Datei → Import und wählen "
        "Sie das gewünschte Format:")
    doc.add_paragraph()

    add_table_2col(doc,
        header=("Format", "Verwendung und Hinweise"),
        rows=[
            ("CSV (ETS6-Export)",
             "ETS6: Extras → Gruppenadressbericht → CSV exportieren. "
             "Importiert alle Gruppenadressen mit Name, Adresse und DPT."),
            ("XLSX (Topologie-Report)",
             "ETS6: Berichte → Topologie-Report → Excel. "
             "Importiert Busstruktur, Linien, Geräte und Kommunikationsobjekte."),
            ("KNXPROJ",
             "Natives ETS-Projektformat. Importiert Topologie und GAs in einem Schritt. "
             "Empfohlenes Format für den vollständigen Projektimport."),
        ]
    )
    doc.add_paragraph()

    add_body(doc,
        "Nach dem Import erscheint das Projekt in der Übersicht. "
        "Die importierten Daten können Sie anschliessend im Wizard oder in den "
        "Einzelansichten (Topologie, Gruppenadressen) prüfen und ergänzen.")
    add_note(doc,
        "Beim Import einer KNXPROJ-Datei werden Gewerk-Zuweisungen nicht übernommen, "
        "da diese ETS-spezifisch sind. Sie können Gewerke in Schritt 5 des Wizards "
        "nachträglich zuweisen.")
    doc.add_paragraph()

    # ── 7  Der Planungs-Wizard ────────────────────────────────────────────────
    add_heading1(doc, "7  Der Planungs-Wizard")
    add_body(doc,
        "Der Wizard führt Sie in 13 Schritten durch die vollständige KNX-Projektplanung. "
        "Starten Sie ihn über Ansicht → Wizard starten (Ctrl+W) oder über die Übersichtsseite. "
        "Sie können jederzeit zwischen den Schritten vor- und zurücknavigieren. "
        "Jeder Schritt zeigt oben rechts einen «?»-Button – dieser öffnet die kontextsensitive "
        "Hilfe direkt zum aktuellen Schritt.")
    doc.add_paragraph()

    add_wizard_step(doc, 1, "Gebäudestruktur",
        "Definieren Sie die Stockwerke Ihres Gebäudes.",
        eingabe=[
            "Stockwerkname (z. B. Kellergeschoss) und Kürzel (z. B. KG)",
            "Hauptgruppen-Nummer (wird automatisch vorgeschlagen)",
            "Alternativ: Schnelleingabe «KG, EG, OG1, DG» für mehrere Stockwerke auf einmal",
        ],
        berechnet=[
            "Hauptgruppen-Nummern für die spätere GA-Generierung",
            "Raumnummernpräfixe pro Stockwerk (E01, O01 usw.)",
        ],
        tipp="Mit «Vorlage laden…» können Sie fertige Gebäudestrukturen (EFH, MFH, Büro) "
             "als Ausgangspunkt verwenden.")

    add_wizard_step(doc, 2, "Wohnungen / Zonen",
        "Unterteilen Sie das Gebäude in logische Zonen oder Wohnungen. "
        "Eine Zone entspricht später genau einer KNX-Linie.",
        eingabe=[
            "Zonenname (z. B. «Wohnung 1», «Erdgeschoss», «Büro Nord»)",
            "Stockwerk-Zuordnung: Welche Stockwerke gehören zu dieser Zone?",
            "Maisonette: Eine Zone kann sich über mehrere Stockwerke erstrecken",
        ],
        berechnet=[
            "Jede Zone erhält später genau eine KNX-Linie in der Topologie",
            "Linienzuordnung aller Räume",
        ],
        tipp="EFH: Klicken Sie auf «1 Zone (EFH)» – das erstellt automatisch eine einzige "
             "Zone für alle Stockwerke.")

    add_wizard_step(doc, 3, "Räume",
        "Fügen Sie alle Räume des Gebäudes hinzu. "
        "Die Räume werden im Baum Zone → Stockwerk → Raum angezeigt.",
        eingabe=[
            "Raumname (z. B. Wohnzimmer, Küche, Büro 1.1)",
            "Stockwerk und Zone (per Baumauswahl)",
            "Masseneingabe: mehrere Räume als Textliste einfügen (ein Raum pro Zeile)",
        ],
        berechnet=[
            "Raumnummern automatisch nach Schema (EFH: E01, O01 / MFH: W1-E01)",
            "Grundlage für alle weiteren Schritte",
        ],
        tipp="Mit Strg+Klick können Sie mehrere Räume gleichzeitig auswählen und "
             "gemeinsam Gewerke zuweisen (nächster Schritt).")

    add_wizard_step(doc, 4, "Elektroverteilungen",
        "Erfassen Sie die Haupt- und Unterverteilungen (HV/UV) des Gebäudes. "
        "Aktoren werden in den Verteilungen installiert, nicht direkt im Raum.",
        eingabe=[
            "Standort der Hauptverteilung (HV)",
            "Standort jeder Unterverteilung (UV) und deren Zuordnung zum Stockwerk/Zone",
        ],
        berechnet=[
            "Zuordnung der Aktoren zu HV/UV",
            "Grundlage für Kabellängenberechnungen und Topologie",
        ])

    add_wizard_step(doc, 5, "Gewerke",
        "Weisen Sie jedem Raum die gewünschten KNX-Funktionen (Gewerke) zu. "
        "Dies ist der zentrale Planungsschritt: aus den Gewerken werden "
        "Aktoren, Sensoren und Gruppenadressen abgeleitet.",
        eingabe=[
            "Gewerk-Code: L (Licht), LD (Dimmbar), J (Jalousie), H (Heizung), "
            "KL (Klimaanlage), LDA (DALI), S (Steckdosen) usw.",
            "Anzahl: wie viele Instanzen dieses Gewerks im Raum? (z. B. J × 2 = 2 Storen)",
            "Vorlagen: Wohnzimmer, Schlafzimmer, Küche – fertige Gewerk-Pakete",
        ],
        berechnet=[
            "Benötigte Aktoren (Typ und Kanalanzahl)",
            "Gruppen-Adressen-Blöcke pro Raum und Gewerk",
            "Sensor-Typ und Tastanzahl",
        ],
        tipp="Mit «Aus Raum erstellen» können Sie die Gewerke eines gut konfigurierten "
             "Raums als wiederverwendbare Vorlage speichern.")

    add_wizard_step(doc, 6, "Gerätekonfiguration",
        "Konfigurieren Sie raumspezifische Einstellungen für spezielle Geräte. "
        "In der Matrix-Ansicht sehen Sie, welches Gewerk auf welcher "
        "Tastereinheit (TE) liegt.",
        eingabe=[
            "Tastereinheiten-Zuweisung: Gewerk → Tastereinheit(en) per Checkbox",
            "Anzahl Tastereinheiten pro Raum (SpinBox, 1–9)",
            "DALI-Gruppen und EVG-Konfiguration (falls LDA-Gewerk vorhanden)",
        ],
        berechnet=[
            "Tastenbedarf pro Sensor (bestimmt Sensor-Typ im nächsten Schritt)",
        ],
        tipp="Ein Gewerk kann auf mehreren Tastereinheiten gleichzeitig liegen "
             "(Mehrfachauswahl per Checkbox).")

    add_wizard_step(doc, 7, "Topologie",
        "KNiX Arranger berechnet die KNX-Topologie automatisch aus den Zonen "
        "und der Gerätezahl. Prüfen Sie Bereiche, Linien und Linienkoppler.",
        eingabe=[
            "Manuelle Anpassung: Räume per Dialog zwischen Linien verschieben",
            "Manuelle Gerät hinzufügen: z. B. IP-Schnittstellen, Server",
        ],
        berechnet=[
            "Bereiche (Areas) und Linien pro Zone",
            "Linienkoppler (LK) und Bereichskoppler (BK) automatisch eingeplant",
            "Speisegerät (SV) pro Linie",
            "Farbcodes: Blau = Koppler, Grün = Speisegerät, Grau = kein Produkt zugewiesen",
        ],
        tipp="Die Topologie-Ansicht in der Seitenleiste zeigt dasselbe Ergebnis "
             "mit erweiterter Navigation.")

    add_wizard_step(doc, 8, "Aktoren",
        "Übersicht der ermittelten Aktoren nach Gerätetyp, Kanal und Raum. "
        "Kontrollieren Sie die automatische Kanalzuweisung.",
        eingabe=[
            "Manuelle Kanalzuweisung bei Bedarf",
            "UV-Zuordnung eines Aktors ändern (Zuweisungs-Dialog)",
        ],
        berechnet=[
            "Schalt-, Dimm-, Jalousie- und Heizungsaktoren aus den Gewerken",
            "Gateways für DALI, Multimedia, Wärmepumpe, E-Ladestation",
            "Kanalzuordnung optimiert nach UV-Standort",
        ])

    add_wizard_step(doc, 9, "Szenen",
        "Definieren Sie Licht- und Jalousie-Szenen vor der Gruppenadress-Generierung. "
        "Szenen werden in Hauptgruppe 0 / Mittelgruppe 5 eingetragen.",
        eingabe=[
            "Szenenname (z. B. Kino, Dinner, Lesen, Abwesenheit)",
            "Szenennummer (1–64, KNX-Standard)",
            "Geltungsbereich: Raum, Zone, Wohnung oder Zentral",
            "Aktionen: Gruppenadresse + Zielwert + optionale Verzögerung",
        ],
        berechnet=[
            "Szenen-GAs automatisch in HG 0 / MG 5 eingeplant",
        ],
        tipp="Vorlagen (Kino, Dinner, Panik usw.) können per Klick als Ausgangspunkt "
             "geladen werden.")

    add_wizard_step(doc, 10, "Gruppenadressen",
        "Die Gruppenadressen werden automatisch nach den KNX Swiss Projektrichtlinien "
        "generiert. Wählen Sie die Mittelgruppen-Variante und starten Sie die Generierung.",
        eingabe=[
            "Variante A: Rückmeldungs-GAs in gleicher Mittelgruppe (kompakt, EFH)",
            "Variante B: Rückmeldungs-GAs in separater MG 6/7 (übersichtlich, MFH)",
            "Manuelle GAs (blau) können ergänzt werden und bleiben bei Neugenerierung erhalten",
        ],
        berechnet=[
            "Vollständige 3-stufige GA-Struktur: HG / MG / UG",
            "HG 0 = Zentraladressen, Szenen, Astro-GAs",
            "HG 1+ = Stockwerke aufsteigend",
            "MG 0=Licht, 1=Jalousie, 2=Heizung, 3=Alarm, 4=Allgemein, 5=Szenen usw.",
            "Datenpunkttypen (DPT) für alle GAs",
        ],
        tipp="Nach der Generierung empfiehlt sich ein Klick auf «Validieren», "
             "um Fehler und Warnungen zu prüfen.")

    add_wizard_step(doc, 11, "Sensor-Ermittlung / Funktionszuordnung",
        "KNiX Arranger ermittelt automatisch, welche Bedienelemente (Tastsensoren, "
        "Raumthermostate, Systemsensoren) in jedem Raum benötigt werden, "
        "und weist ihnen Funktionen zu.",
        eingabe=[
            "Sensor-Typ manuell überschreiben (Doppelklick auf Raum)",
            "Funktionszuordnung anpassen: welche Taste steuert welches Gewerk",
        ],
        berechnet=[
            "Sensor-Typ und Tastenanzahl basierend auf Gewerken und Tastereinheiten",
            "Tastsensor 2-fach / 4-fach / 8-fach je nach Gewerk-Anzahl",
            "Raumthermostat wenn Heizungs- oder Klimagewerk vorhanden",
            "Systemsensoren (Wetterstation, Windsensor) aus Zentral-Gewerken",
        ])

    add_wizard_step(doc, 12, "Funktionsdefinition",
        "Das Bauherr-Formular zeigt alle Bedienelemente und deren zugeordneten Funktionen. "
        "Es dient als Grundlage für die Inbetriebnahme und wird vom Bauherrn unterschrieben.",
        eingabe=[
            "Tastenbelegung prüfen und bei Bedarf manuell anpassen",
            "Beschreibungstext pro Bedienelement ergänzen",
        ],
        berechnet=[
            "Bauherr-Formular als XLSX exportierbar (Schritt 13)",
            "Rückmelde-GA wird automatisch zum Steuertelegramm ergänzt (Status-LED)",
        ],
        tipp="«Auto-Assign» belegt alle Tasten automatisch. "
             "Danach können Einzeltasten manuell angepasst werden.")

    add_wizard_step(doc, 13, "Export",
        "Exportieren Sie das Projekt in die gewünschten Formate.",
        eingabe=[
            "Zielordner für das Revisionspaket wählen",
            "Exportformat auswählen",
        ],
        berechnet=None)

    add_table_2col(doc,
        header=("Exportformat", "Inhalt / Verwendung"),
        rows=[
            ("CSV",
             "ETS6-kompatibler Gruppenadress-Export. "
             "Direkt in ETS6 importierbar (Extras → Gruppenadressbericht → Import)."),
            ("KNXPROJ",
             "Natives ETS6-Projektformat mit Topologie und Gruppenadressen. "
             "Vollständiger Import in ETS6 möglich."),
            ("Revisionspaket (ZIP)",
             "Vollständiges Dokumentationspaket: GA-Übersicht (XLSX), Topologie-Plan, "
             "Belegungsplan, Kabellängen-Report, Datenbläter."),
            ("Bauherr-Formular",
             "Tastenbelegungen als XLSX – wird vom Bauherrn unterschrieben."),
            ("Abnahmeprotokoll",
             "Vordruck für die Projektübergabe und Inbetriebnahme."),
        ]
    )
    doc.add_paragraph()

    # ── 8  Wichtige Ansichten ─────────────────────────────────────────────────
    add_heading1(doc, "8  Wichtige Ansichten in der Seitenleiste")
    add_body(doc,
        "Alle Ansichten sind auch ausserhalb des Wizards jederzeit zugänglich. "
        "So können Sie z. B. nach Abschluss des Wizards einzelne Bereiche direkt bearbeiten.")
    doc.add_paragraph()

    add_table_2col(doc,
        header=("Ansicht", "Wann verwenden?"),
        rows=[
            ("Übersicht",
             "Einstieg: Projektfortschritt, Schnellzugriff auf Wizard und Export."),
            ("Gebäude",
             "Baumstruktur: Gebäude → Flügel → Stockwerk → Zone → Raum. "
             "Für manuelle Anpassungen an der Gebäudestruktur nach dem Wizard."),
            ("Topologie",
             "KNX-Linienbaum mit Kopplern und Speisegeräten. "
             "Für manuelle Gerätezuweisungen und Leistungsberechnungen."),
            ("Gruppenadressen",
             "Tabellarische GA-Übersicht mit Filterfunktion. "
             "Für manuelle GA-Korrekturen und GA-Suche."),
            ("Verknüpfungsmatrix",
             "Welcher Sensor sendet auf welche GA? "
             "Für die Kontrolle vor der ETS-Programmierung."),
            ("Materialliste",
             "Ermittelte Geräte mit Typ, Kanal und Raum. "
             "Für Produktzuweisung und Offertanfragen."),
            ("Szenen",
             "Szenen-Verwaltung ausserhalb des Wizards. "
             "Für Nachträge oder Korrekturen an Szenen."),
            ("Zeitsteuerung",
             "Wochenprogramme mit Astrofunktionen und Feiertagskalender. "
             "Für zeitgesteuerte Schaltprogramme."),
            ("DALI-Konfiguration",
             "DALI-Gruppen und EVGs – automatisch aus importierten GAs ermittelt. "
             "Nur relevant wenn DALI-Leuchten (LDA-Gewerk) vorhanden."),
            ("Validierung",
             "Prüft das Projekt auf Fehler und Warnungen vor dem Export. "
             "Empfohlen immer vor dem CSV/KNXPROJ-Export."),
            ("Export",
             "Direktzugriff auf alle Exportfunktionen (auch ohne Wizard)."),
            ("Gewerk-Katalog",
             "Übersicht aller verfügbaren Gewerk-Codes mit Beschreibung. "
             "Zum Nachschlagen während der Gewerke-Zuweisung (Schritt 5)."),
            ("Bauherren-Beratung",
             "Interaktive Tastenbelegung gemeinsam mit dem Bauherrn: Für jede Taste "
             "kann ein Wunsch ausgewählt und je Raum eine Anmerkung erfasst werden. "
             "Ideal für die Besprechung am Bildschirm vor dem Bauherr-Formular (Schritt 12)."),
            ("Datenblätter",
             "Produktdatenblätter zu den zugewiesenen Geräten, inkl. Online-Suche "
             "und Speichern von Links."),
            ("Kundenofferten",
             "Automatische Angebotskalkulation basierend auf Materialliste und den "
             "Stundensätzen aus dem Firmenprofil."),
            ("KNX Secure",
             "Konfiguration der KNX-Secure-Schlüssel für gesicherte Linien und Geräte."),
            ("Inbetriebnahme",
             "Checkliste und Protokoll für die Projektübergabe "
             "(siehe auch Abnahmeprotokoll, Schritt 13)."),
        ]
    )
    doc.add_paragraph()

    # ── 9  Tastenkürzel ───────────────────────────────────────────────────────
    add_heading1(doc, "9  Wichtige Tastenkürzel")

    add_table_2col(doc, [
        ("Ctrl+N",       "Neues Projekt"),
        ("Ctrl+O",       "Projekt öffnen"),
        ("Ctrl+S",       "Speichern"),
        ("Ctrl+Shift+S", "Speichern unter"),
        ("Ctrl+W",       "Wizard starten"),
        ("Ctrl+I",       "Import (CSV / XLSX / KNXPROJ)"),
        ("Ctrl+E",       "CSV-Export"),
        ("Ctrl+Z",       "Rückgängig"),
        ("Ctrl+Y",       "Wiederholen"),
        ("Ctrl+V",       "Validieren"),
        ("F1",           "Hilfe anzeigen (kontextsensitiv)"),
        ("Esc",          "Dialog / Wizard schliessen"),
    ])
    doc.add_paragraph()

    # ── 10  Hilfe-System ──────────────────────────────────────────────────────
    add_heading1(doc, "10  Hilfe und Onboarding")

    add_heading2(doc, "10.1  Kontextsensitive Hilfe (F1 / «?»-Button)")
    add_body(doc,
        "Drücken Sie jederzeit F1, um die Hilfe zur aktuell geöffneten Ansicht anzuzeigen. "
        "In jedem Wizard-Schritt finden Sie oben rechts einen «?»-Button – dieser öffnet "
        "direkt das passende Hilfethema.")
    add_body(doc,
        "Die Hilfe-Ansicht (Seitenleiste → Hilfe) enthält alle Themen als durchsuchbare Liste. "
        "Geben Sie einen Suchbegriff ein, um relevante Hilfeseiten zu filtern.")
    doc.add_paragraph()

    add_heading2(doc, "10.2  Onboarding-Tour")
    add_body(doc,
        "Beim ersten Start des Programms erscheint automatisch die Onboarding-Tour. "
        "Sie führt in 6 Schritten durch die wichtigsten Konzepte (Gewerke, Topologie, "
        "Gruppenadressen, Wizard-Ablauf).")
    add_bullet(doc,
        "Tour erneut starten: Hilfe → Erste Schritte")
    add_bullet(doc,
        "Tour überspringen: «Tour überspringen»-Button im Dialog")
    add_tip(doc,
        "Die Tour kann über die Checkbox «Tour beim nächsten Start nicht mehr anzeigen» "
        "dauerhaft deaktiviert werden.")
    doc.add_paragraph()

    # ── 11  Troubleshooting ───────────────────────────────────────────────────
    add_heading1(doc, "11  Häufige Probleme (Troubleshooting)")

    add_table_2col(doc,
        header=("Problem", "Lösung"),
        rows=[
            ("Windows SmartScreen blockiert den Start",
             "«Weitere Informationen» → «Trotzdem ausführen». "
             "Das Programm ist nicht signiert, aber sicher."),
            ("Lizenzdialog erscheint beim Start",
             "Die .knxlic-Datei wurde noch nicht importiert oder ist abgelaufen. "
             "Klicken Sie auf «Lizenzdatei auswählen» und wählen Sie Ihre .knxlic-Datei."),
            ("«Lizenz ungültig» Fehlermeldung",
             "Prüfen Sie das Ablaufdatum der Datei (steht im Dateinamen). "
             "Melden Sie sich bei uns für eine Verlängerung."),
            ("Export schlägt fehl / leere Datei",
             "Stellen Sie sicher, dass das Projekt validiert wurde (Ctrl+V) und "
             "keine Fehler vorhanden sind. Prüfen Sie den Zielordner auf Schreibrechte."),
            ("Gruppenadressen werden nicht generiert",
             "Prüfen Sie ob Gewerke in Schritt 5 zugewiesen wurden und "
             "ob die Topologie (Schritt 7) berechnet wurde."),
            ("Programm startet nicht / Absturz",
             "Log-Datei prüfen: %APPDATA%\\KNiX Arranger\\knix_arranger.log. "
             "Bitte legen Sie diese Datei dem Feedback-Mail bei."),
            ("KNXPROJ-Import schlägt fehl",
             "Stellen Sie sicher, dass die Datei aus ETS6 exportiert wurde "
             "(nicht ETS5 oder älter). Bei Fehlern: knxproj-Datei dem Feedback beilegen."),
            ("Anzeigeprobleme / UI abgeschnitten",
             "Mindestauflösung 1280 × 768. Prüfen Sie die Windows-Skalierung "
             "(empfohlen: 100 % bei 1920 × 1080)."),
        ]
    )
    doc.add_paragraph()

    # ── 12  Bekannte Einschränkungen ──────────────────────────────────────────
    add_heading1(doc, "12  Bekannte Einschränkungen (Beta-Version)")
    add_body(doc,
        "Als Betatester erhalten Sie eine Vorabversion. Folgende Punkte sind bekannt:")
    add_bullet(doc, "Einzelne Exportfunktionen können noch unvollständig sein.")
    add_bullet(doc, "Die integrierte Hilfedokumentation wird noch ausgebaut.")
    add_bullet(doc, "Das Programm wurde noch nicht auf allen Systemkonfigurationen getestet.")
    add_bullet(doc, "Screenshots in der Dokumentation folgen in der nächsten Version.")
    doc.add_paragraph()

    # ── 13  Feedback ─────────────────────────────────────────────────────────
    add_heading1(doc, "13  Feedback erwünscht!")
    add_body(doc,
        "Ihr Feedback ist für uns sehr wertvoll. Bitte melden Sie uns:")
    add_bullet(doc, "Fehler und unerwartetes Verhalten (mit kurzer Beschreibung der Schritte)")
    add_bullet(doc, "Unklare oder fehlende Funktionen")
    add_bullet(doc, "Verbesserungsvorschläge zur Bedienbarkeit")
    add_bullet(doc, "Allgemeine Eindrücke zur Anwendung")
    doc.add_paragraph()
    add_body(doc,
        "Feedback per E-Mail an:  info@muellersmarthomeenergiemanagement.ch")
    add_body(doc,
        "Bitte geben Sie in der Betreffzeile «KNiX Arranger Beta» an. "
        "Für Absturz-Reports finden Sie unter %APPDATA%\\KNiX Arranger\\ eine "
        "Log-Datei, die Sie uns beilegen können.")
    doc.add_paragraph()

    add_heading2(doc, "13.1  Feedback-Vorlage")
    add_body(doc,
        "Um uns die Bearbeitung zu erleichtern, verwenden Sie bitte diese Vorlage "
        "als Grundlage für Ihr Feedback-Mail:")
    doc.add_paragraph()

    add_table_2col(doc,
        header=("Feld", "Ihr Inhalt"),
        rows=[
            ("Betreff",          "KNiX Arranger Beta – [Art: Fehler / Verbesserung / Frage]"),
            ("Programmversion",  "Zu finden unter Hilfe → Über KNiX Arranger"),
            ("Windows-Version",  "z. B. Windows 11 Pro 23H2"),
            ("Schritt / Ansicht",
             "Wo trat das Problem auf? (z. B. Wizard Schritt 5, Topologie-Ansicht)"),
            ("Was wurde gemacht",
             "Kurze Schrittfolge: z. B. «Schritt 5 → Vorlage Wohnzimmer angewählt "
             "→ Auf Raum anwenden → Absturz»"),
            ("Was wurde erwartet", "Was sollte eigentlich passieren?"),
            ("Was ist passiert",   "Was ist tatsächlich passiert? Fehlermeldung?"),
            ("Log-Datei",
             "Bitte beilegen: %APPDATA%\\KNiX Arranger\\knix_arranger.log"),
            ("Projektdatei",
             "Falls möglich: .knxarr-Datei beilegen (anonymisiert falls nötig)"),
        ]
    )
    doc.add_paragraph()

    # ── Abschluss ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Herzlichen Dank für Ihre Unterstützung!")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_GREEN_DARK

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mueller SmartHome & EnergieManagement")
    run.font.size = Pt(11)

    return doc


if __name__ == "__main__":
    import sys, io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    doc = build_document()
    out_path = "KNiX_Arranger_Betatester_Anleitung.docx"
    doc.save(out_path)
    print(f"Gespeichert: {out_path}")
    print(f"Absätze: {len(doc.paragraphs)}, Tabellen: {len(doc.tables)}")
